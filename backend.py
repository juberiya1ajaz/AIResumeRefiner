# Imports for file handling, NLP, PDF/Word processing, API requests, and temporary file creation
import os
import re
import requests
import nltk
import PyPDF2
import docx
import spacy
import streamlit as st

from io import BytesIO
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit

# ---------------------------
# NLTK & SpaCy Setup
# ---------------------------

nltk_data_path = os.path.join(os.path.dirname(__file__), "nltk_data")
os.makedirs(nltk_data_path, exist_ok=True)
nltk.data.path.append(nltk_data_path)

nltk.download("punkt", download_dir=nltk_data_path)
nltk.download("stopwords", download_dir=nltk_data_path)

@st.cache_resource
def get_nlp():
    return spacy.load("en_core_web_sm")

nlp = get_nlp()
stop_words = set(stopwords.words("english"))

# ---------------------------
# PERPLEXITY API Key
# ---------------------------

PERPLEXITY_API_KEY = st.secrets["PERPLEXITY_API_KEY"]

# ---------------------------
# TEXT EXTRACTION FUNCTIONS
# ---------------------------

def extract_text(resume_file):
    if resume_file.name.endswith(".pdf"):
        reader = PyPDF2.PdfReader(resume_file)
        return "".join([page.extract_text() for page in reader.pages])
    elif resume_file.name.endswith(".docx"):
        doc = docx.Document(resume_file)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        return "Unsupported file type."

# ---------------------------
# KEYWORD & SKILL MATCHING
# ---------------------------

def extract_keywords(text, top_n=10):
    tokens = word_tokenize(text)
    words = [w.lower() for w in tokens if w.isalpha() and w.lower() not in stop_words]
    freq_dist = nltk.FreqDist(words)
    return [word for word, freq in freq_dist.most_common(top_n)]

def ats_keyword_check(resume_text, job_desc):
    resume_keywords = set(extract_keywords(resume_text, 50))
    jd_keywords = set(extract_keywords(job_desc, 50))
    matching = resume_keywords & jd_keywords
    coverage = len(matching) / (len(jd_keywords) or 1)
    suggestions = jd_keywords - resume_keywords
    return {
        "matching_keywords": list(matching),
        "missing_keywords": list(suggestions),
        "coverage_percent": round(coverage * 100, 2)
    }

# ---------------------------
# FILE EXPORTING
# ---------------------------

def export_to_docx(text):
    buffer = BytesIO()
    doc = Document()
    bold_pattern = re.compile(r'\*\*(.*?)\*\*')

    for line in text.split('\n'):
        para = doc.add_paragraph()
        if line.startswith("[Score & Feedback]"):
            run = para.add_run("Score & Feedback:\n")
            run.bold = True
            continue

        pos = 0
        for match in bold_pattern.finditer(line):
            start, end = match.span()
            if start > pos:
                para.add_run(line[pos:start])
            bold_run = para.add_run(match.group(1))
            bold_run.bold = True
            pos = end
        if pos < len(line):
            para.add_run(line[pos:])
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_to_pdf(text):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    x_margin = 40
    y = height - 40
    font_size = 11
    normal_font = "Helvetica"
    bold_font = "Helvetica-Bold"
    line_spacing = 10

    def draw_wrapped_line(parts):
        nonlocal y
        x = x_margin
        for part, is_bold in parts:
            font = bold_font if is_bold else normal_font
            c.setFont(font, font_size)
            for word in part.split():
                word_width = c.stringWidth(word + ' ', font, font_size)
                if x + word_width > width - x_margin:
                    y -= line_spacing
                    x = x_margin
                    if y < 40:
                        c.showPage()
                        y = height - 40
                        c.setFont(font, font_size)
                c.drawString(x, y, word + ' ')
                x += word_width
        y -= line_spacing
        if y < 40:
            c.showPage()
            y = height - 40

    bold_pattern = re.compile(r'\*\*(.*?)\*\*')

    for line in text.strip().split('\n'):
        line = line.strip()
        if not line:
            y -= 10
            continue
        if line.startswith("[Score & Feedback]"):
            c.setFont(bold_font, font_size)
            c.drawString(x_margin, y, "Score & Feedback:")
            y -= line_spacing
            continue

        parts = []
        last = 0
        for match in bold_pattern.finditer(line):
            if match.start() > last:
                parts.append((line[last:match.start()], False))
            parts.append((match.group(1), True))
            last = match.end()
        if last < len(line):
            parts.append((line[last:], False))

        draw_wrapped_line(parts)

    c.save()
    buffer.seek(0)
    return buffer

# ---------------------------
# PERPLEXITY API INTEGRATION
# ---------------------------

def call_perplexity(prompt):
    url = "https://api.perplexity.ai/chat/completions"
    headers = {
        "Authorization": f"Bearer {PERPLEXITY_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "sonar-pro",
        "messages": [{"role": "user", "content": prompt}]
    }
    response = requests.post(url, headers=headers, json=data)
    response.raise_for_status()
    return response.json()['choices'][0]['message']['content']

# ---------------------------
# SKILL EXTRACTION
# ---------------------------

def fetch_dynamic_skillset_from_perplexity(job_desc):
    prompt = (
        "Extract a list of up to 50 relevant technical and soft skills from the following job description. "
        "Only return a comma-separated list of skill names, no explanations:\n\n" + job_desc
    )
    skills_raw = call_perplexity(prompt)
    skills = re.split(r',\s*|\n', skills_raw)
    return set(skill.strip().lower() for skill in skills if skill.strip())

def extract_skills(text, skillset):
    doc = nlp(text)
    return set(token.text.lower() for token in doc if token.text.lower() in skillset)

# ---------------------------
# STRUCTURED SECTION PARSING
# ---------------------------

def parse_sections_with_llm(resume_text):
    prompt = (
        "You are an expert resume parser. Given the following unstructured resume text, extract and organize it into a clean structured format. "
        "Detect all relevant sections, even if they are not standard (like Certifications, Projects, Publications, Languages, etc). "
        "Return each section with a clear heading like this:\n\n"
        "=== Section Name ===\n"
        "Section content...\n\n"
        f"Resume:\n{resume_text}"
    )
    return call_perplexity(prompt)

# ---------------------------
# SECTION SCORING
# ---------------------------

def score_section_with_llm(section_name, section_content):
    prompt = (
        f"Evaluate the following '{section_name}' section of a resume. "
        "Score it from 1 to 10 based on clarity, impact, and relevance to a typical job description. "
        "Also give 1-2 sentences of constructive feedback.\n\n"
        f"Section Content:\n{section_content}"
    )
    return call_perplexity(prompt)

def score_all_sections(parsed_resume):
    sections = re.findall(r'=== (.*?) ===\n(.*?)(?=(?:===|\Z))', parsed_resume, re.DOTALL)
    section_scores = {}
    for title, content in sections:
        result = score_section_with_llm(title.strip(), content.strip())
        section_scores[title.strip()] = result
    return section_scores

def regenerate_section_with_llm(section_name, current_content):
    prompt = (
        f"You are an expert resume writer. Rewrite the '{section_name}' section below to improve clarity, impact, and alignment with industry best practices. "
        "Keep it concise and relevant. Return only the improved content without section headers or additional commentary.\n\n"
        f"{section_name} Section:\n{current_content}"
    )
    return call_perplexity(prompt)

# ---------------------------
# RESUME OPTIMIZATION
# ---------------------------

def optimize_resume(resume_text, job_desc, skillset):
    prompt = (
        "You are an expert resume optimizer. Given the following resume and job description, tailor the resume towards the job, "
        "add important missing keywords and skills, and suggest ATS improvements. Highlight additions and changes.\n\n"
        f"Relevant Skills: {', '.join(skillset)}\n\n"
        f"Resume:\n{resume_text}\n\nJob Description:\n{job_desc}"
    )
    optimized_resume = call_perplexity(prompt)
    resume_skills = extract_skills(resume_text.lower(), skillset)
    jd_skills = extract_skills(job_desc.lower(), skillset)
    return optimized_resume, resume_skills, jd_skills
